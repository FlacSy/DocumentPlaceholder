"""Read a Word template, replace ``{PLACEHOLDER}`` markers, and save."""

from __future__ import annotations

import urllib.request
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Cm

from document_placeholder.image_value import ImageValue


class DocumentProcessor:
    def __init__(self, template_path: str | Path) -> None:
        self.doc = Document(str(template_path))

    # -- public API -----------------------------------------------------------

    def replace_placeholders(self, values: dict[str, Any]) -> None:
        """Substitute every ``{KEY}`` found in paragraphs, tables, headers, and footers."""

        for paragraph in self.doc.paragraphs:
            self._replace_in_paragraph(paragraph, values)

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, values)

        for section in self.doc.sections:
            for paragraph in section.header.paragraphs:
                self._replace_in_paragraph(paragraph, values)
            for paragraph in section.footer.paragraphs:
                self._replace_in_paragraph(paragraph, values)

    def save(self, output_path: str | Path) -> None:
        self.doc.save(str(output_path))

    # -- internals ------------------------------------------------------------

    @staticmethod
    def _replace_in_paragraph(paragraph, values: dict[str, Any]) -> None:
        runs = paragraph.runs
        if not runs:
            return

        for key, value in values.items():
            if not isinstance(value, ImageValue):
                continue
            placeholder = "{" + key + "}"
            for run in runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, "")
                    try:
                        stream = DocumentProcessor._load_image(value.source)
                        w = value.width_cm if value.width_cm is not None else 5.0
                        kwargs: dict[str, Any] = {"width": Cm(w)}
                        if value.height_cm is not None:
                            kwargs["height"] = Cm(value.height_cm)
                        run.add_picture(stream, **kwargs)
                    except (OSError, ValueError, KeyError):
                        pass
                    break

        full_text = "".join(run.text for run in runs)
        new_text = full_text
        for key, value in values.items():
            if isinstance(value, ImageValue):
                continue
            placeholder = "{" + key + "}"
            if placeholder in new_text:
                display = str(value) if value is not None else ""
                display = DocumentProcessor._sanitize_xml_text(display)
                new_text = new_text.replace(placeholder, display)

        if new_text != full_text:
            runs[0].text = DocumentProcessor._sanitize_xml_text(new_text)
            for run in runs[1:]:
                run.text = ""

    @staticmethod
    def _sanitize_xml_text(text: str) -> str:
        """Удалить символы, недопустимые в XML (NULL, control chars)."""
        if not text:
            return ""
        result = []
        for c in str(text):
            code = ord(c)
            if code == 0x9 or code == 0xA or code == 0xD:
                result.append(c)
            elif 0x20 <= code <= 0xD7FF or 0xE000 <= code <= 0xFFFD:
                result.append(c)
            elif 0x10000 <= code <= 0x10FFFF:
                result.append(c)
            else:
                result.append(" ")
        return "".join(result)

    @staticmethod
    def _load_image(source: str) -> BytesIO:
        """Загрузить изображение из URL или файла. SVG конвертируется в PNG."""
        source = source.strip()
        if source.startswith(("http://", "https://")):
            req = urllib.request.Request(
                source,
                headers={"User-Agent": "DocumentPlaceholder/1.0"},
            )
            with urllib.request.urlopen(req, timeout=30) as resp:
                data = resp.read()
        else:
            path = Path(source)
            if path.exists():
                data = path.read_bytes()
            else:
                raise FileNotFoundError(source)

        return DocumentProcessor._ensure_raster(data, source)

    @staticmethod
    def _ensure_raster(data: bytes, source: str = "") -> BytesIO:
        """Конвертировать SVG в PNG, остальные форматы вернуть как есть."""
        if DocumentProcessor._is_svg(data, source):
            try:
                from cairosvg import svg2png

                png_out = BytesIO()
                svg2png(bytestring=data, write_to=png_out)
                png_out.seek(0)
                return png_out
            except Exception:
                raise ValueError("SVG conversion failed")
        return BytesIO(data)

    @staticmethod
    def _is_svg(data: bytes, source: str) -> bool:
        """Проверить, является ли контент SVG."""
        if source.lower().endswith(".svg") or source.lower().endswith(".svgz"):
            return True
        start = data.lstrip()[:200].decode("utf-8", errors="ignore")
        return start.lstrip().startswith("<svg") or start.lstrip().startswith("<?xml")
